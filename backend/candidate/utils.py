import os
import io
import subprocess
import warnings

# Suppress spaCy warnings that might interfere with parsing
warnings.filterwarnings('ignore', category=UserWarning, module='spacy')

# Lazy import for resume parser to avoid NLTK loading during Django startup
# This prevents NLTK stopwords errors during migrations
ResumeParser = None
resume_parser_available = None  # Will be checked when needed
from django.core.files.uploadedfile import InMemoryUploadedFile

# NOTE: The docx2pdf library requires LibreOffice to be installed and accessible
# on your system path.
from docx2pdf import convert

def _lazy_load_resume_parser():
    """Lazy load ResumeParser only when needed to avoid NLTK issues during startup"""
    global ResumeParser, resume_parser_available
    
    if resume_parser_available is None:
        try:
            from pyresparser import ResumeParser
            resume_parser_available = True
            print(" ResumeParser loaded successfully")
        except ImportError as e:
            ResumeParser = None
            resume_parser_available = False
            print(f" ResumeParser not available: {e}")
        except Exception as e:
            ResumeParser = None
            resume_parser_available = False
            print(f" Error loading ResumeParser: {e}")
    
    return resume_parser_available

def get_or_create_source(source_name):
    """Get or create a Source object"""
    if not source_name:
        return None
    
    from .models import Source
    source, created = Source.objects.get_or_create(source_name=source_name)
    return source.id

def get_or_create_designation(designation_name):
    """Get or create a Designation object"""
    if not designation_name:
        return None
    
    from .models import Designation
    designation, created = Designation.objects.get_or_create(designation_name=designation_name)
    return designation.id

def clean_experience_data(experience_value):
    """Clean and convert experience data to proper decimal format"""
    if not experience_value:
        return 0.0
    
    try:
        # Convert to string first
        exp_str = str(experience_value).strip()
        
        # Remove any non-numeric characters except decimal point
        import re
        exp_clean = re.sub(r'[^\d.]', '', exp_str)
        
        # Convert to float
        exp_float = float(exp_clean) if exp_clean else 0.0
        
        # Ensure it's within valid range for DecimalField
        if exp_float > 99.9:
            exp_float = 99.9
        elif exp_float < 0:
            exp_float = 0.0
            
        return exp_float
    except (ValueError, TypeError):
        return 0.0

def parse_resume(file):
    """
    Parses resume data using the resume_parser library.
    This function handles both in-memory uploaded files and file paths.
    """
    if not file:
        return {"success": False, "error": "No file provided."}

    # If the file is an in-memory object, save it to a temporary file for the parser.
    if isinstance(file, InMemoryUploadedFile):
        temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Generate a unique filename to avoid conflicts
        unique_filename = f"{os.urandom(8).hex()}_{file.name}"
        temp_file_path = os.path.join(temp_dir, unique_filename)
        
        with open(temp_file_path, 'wb+') as temp_file:
            for chunk in file.chunks():
                temp_file.write(chunk)
    else:
        # If 'file' is already a path string, use it directly.
        temp_file_path = file

    # Convert .docx to .pdf if necessary.
    # The resumeparse library may handle .docx directly, but converting to PDF is a
    # robust fallback and often improves accuracy.
    if temp_file_path.lower().endswith('.docx'):
        try:
            pdf_path = convert_docx_to_pdf(temp_file_path, temp_dir)
            if pdf_path:  # Only use PDF path if conversion was successful
                temp_file_path = pdf_path
        except Exception as e:
            # Handle conversion failure
            print(f"Error converting DOCX to PDF: {e}")
            # The function can still try to parse the docx file directly
            pass

    try:
        # Lazy load ResumeParser only when actually needed
        if not _lazy_load_resume_parser():
            print("[WARNING] pyresparser not available, using alternative parser...")
            # Don't return error, fall through to alternative parser
        
        # Skip pyresparser entirely and use alternative parser for E985 errors
        # This is a more reliable approach given the spaCy configuration issues
        try:
            from .alternative_parser import alternative_parse_resume
            print("Using alternative parser to avoid spaCy configuration issues...")
            fallback_result = alternative_parse_resume(temp_file_path)
            if fallback_result["success"]:
                fallback_result["message"] = "Used alternative parser (recommended for stability)"
                return fallback_result
        except Exception as alt_error:
            print(f"Alternative parser failed: {alt_error}")
        
        # If alternative parser fails, try simple text extraction
        try:
            fallback_data = simple_text_extraction(temp_file_path)
            if fallback_data:
                return {
                    "success": True,
                    "data": fallback_data,
                    "message": "Used simple text extraction method"
                }
        except Exception as simple_error:
            print(f"Simple text extraction failed: {simple_error}")
        
        # Only try pyresparser as a last resort (keeping original code for reference)
        try:
            # Initialize spaCy with better error handling
            import spacy
            nlp = None
            
            # Try multiple approaches to load spaCy model
            model_names = ["en_core_web_sm", "en_core_web_md", "en_core_web_lg"]
            for model_name in model_names:
                try:
                    nlp = spacy.load(model_name)
                    break
                except OSError:
                    continue
            
            # If no model found, create a blank English model
            if nlp is None:
                try:
                    nlp = spacy.blank("en")
                    print("Warning: Using blank spaCy model. Consider installing en_core_web_sm for better results.")
                except Exception as spacy_error:
                    print(f"Warning: Could not initialize spaCy: {spacy_error}")
                    # Skip pyresparser if spaCy fails
                    raise Exception("spaCy initialization failed")
            
            # Create pyresparser config directory and file if needed
            try:
                import pkg_resources
                pyresparser_path = pkg_resources.get_distribution("pyresparser").location
                config_dir = os.path.join(pyresparser_path, "pyresparser")
                config_path = os.path.join(config_dir, "config.cfg")
                
                if not os.path.exists(config_path):
                    os.makedirs(config_dir, exist_ok=True)
                    with open(config_path, 'w') as f:
                        f.write("[nlp]\n")
                        f.write("lang = en\n")
                        f.write("pipeline = tagger,parser,ner\n")
                        f.write("\n[DEFAULT]\n")
                        f.write("# Basic pyresparser configuration\n")
            except Exception as config_error:
                print(f"Warning: Could not create config file: {config_error}")
                # Skip pyresparser if config creation fails
                raise Exception("Config file creation failed")
            
            # Try to initialize ResumeParser with error handling
            parser = ResumeParser(temp_file_path)
            data = parser.get_extracted_data()
            
        except Exception as e:
            error_msg = str(e)
            print(f"PyResParser error: {error_msg}")
            
            # Try fallback parsing method for common errors
            if any(keyword in error_msg.lower() for keyword in ["config", "e053", "spacy", "e1005", "pos", "nlp"]):
                try:
                    # Import and use alternative parser
                    from .alternative_parser import alternative_parse_resume
                    fallback_result = alternative_parse_resume(temp_file_path)
                    if fallback_result["success"]:
                        fallback_result["message"] = "Used alternative parser due to spaCy/pyresparser configuration issues"
                        return fallback_result
                except Exception as fallback_error:
                    print(f"Alternative parser also failed: {fallback_error}")
                
                # If alternative parser fails, try simple text extraction
                try:
                    fallback_data = simple_text_extraction(temp_file_path)
                    if fallback_data:
                        return {
                            "success": True,
                            "data": fallback_data,
                            "message": "Used simple text extraction as final fallback"
                        }
                except Exception as simple_fallback_error:
                    print(f"Simple text extraction also failed: {simple_fallback_error}")
            
            # Return the original error if all fallbacks fail
            return {"success": False, "error": f"Failed to parse resume: {error_msg}"}
        
        # Handle different data formats
        if isinstance(data, dict):
            parsed_data = {
                "name": data.get("name") or data.get("Name"),
                "email": data.get("email") or data.get("Email"),
                "phone": data.get("mobile_number") or data.get("phone") or data.get("Phone"),
                "mobile": data.get("mobile_number") or data.get("mobile") or data.get("Mobile"),
                "skills": data.get("skills") or data.get("Skills"),
                "language": data.get("language") or data.get("Language"),
                "education": {
                    "college_name": data.get("college_name") or data.get("College Name"),
                    "degree": data.get("degree") or data.get("Degree"),
                },
                "experience": {
                    "designation": data.get("designation") or data.get("Designation"),
                    "company_names": data.get("company_names") or data.get("Company Names"),
                    "total_experience": data.get("total_experience") or data.get("Total Experience", 0),
                },
                "no_of_pages": data.get("no_of_pages") or data.get("No of Pages"),
                "raw_text": str(data) if not data.get("raw_text") else data.get("raw_text"),
            }
            return {"success": True, "data": parsed_data}
        else:
            # If data is not a dict, try to extract basic info
            return {
                "success": True, 
                "data": {
                    "raw_text": str(data),
                    "message": "Basic parsing completed. Manual review recommended."
                }
            }

    except Exception as e:
        return {"success": False, "error": f"Error during resume parsing: {e}"}

    finally:
        # Clean up all temporary files created by this function.
        try:
            if isinstance(file, InMemoryUploadedFile):
                # Remove the original temp file
                if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up temporary files: {cleanup_error}")

def simple_text_extraction(file_path):
    """
    Simple fallback text extraction for resume parsing when pyresparser fails.
    Extracts basic information using simple text processing.
    """
    import re
    
    try:
        text_content = ""
        
        # Extract text based on file type
        if file_path.lower().endswith('.pdf'):
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
            except ImportError:
                # If PyPDF2 is not available, try pdfplumber
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            text_content += page.extract_text() + "\n"
                except ImportError:
                    return None
        
        elif file_path.lower().endswith(('.doc', '.docx')):
            try:
                from docx import Document
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
            except ImportError:
                try:
                    import python_docx
                    from python_docx import Document
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except ImportError:
                    print("Warning: python-docx not available for DOCX parsing")
                    return None
        
        else:
            # Try to read as plain text
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text_content = file.read()
            except:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
        
        if not text_content.strip():
            return None
        
        # Simple pattern matching for basic information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        
        # Improved phone patterns for Indian numbers and international formats
        phone_patterns = [
            r'\+91\s*\d{10}',  # +91 followed by 10 digits
            r'\+91[-.\s]\d{10}',  # +91 with separator and 10 digits
            r'\+91\s*\d{5}\s*\d{5}',  # +91 with space in middle
            r'\b\d{10}\b',  # 10 digit Indian mobile numbers
            r'\(\+91\)\s*\d{10}',  # (+91) format
            r'\+\d{1,3}[-.\s]?\d{6,14}',  # General international format
        ]
        
        emails = re.findall(email_pattern, text_content)
        
        # Try each phone pattern and collect all matches
        phones = []
        for pattern in phone_patterns:
            matches = re.findall(pattern, text_content)
            phones.extend(matches)
        
        # Clean up phone numbers (remove extra spaces)
        phones = [re.sub(r'\s+', ' ', phone.strip()) for phone in phones if phone.strip()]
        
        # Try to extract name from first line or common patterns
        lines = text_content.strip().split('\n')
        potential_name = ""
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if len(line) > 2 and len(line) < 50 and ' ' in line:
                # Simple heuristic: if line has 2-4 words and looks like a name
                words = line.split()
                if 2 <= len(words) <= 4 and all(word.replace('.', '').isalpha() for word in words):
                    potential_name = line
                    break
        
        # Extract Education Information
        education_info = extract_education(text_content)
        
        # Extract Experience Information  
        experience_info = extract_experience(text_content)
        
        # Extract Languages
        languages = extract_languages(text_content)
        
        # Extract skills (simple keyword matching)
        skill_keywords = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'django', 'flask',
            'html', 'css', 'sql', 'mysql', 'postgresql', 'mongodb', 'aws', 'azure', 'docker',
            'kubernetes', 'git', 'agile', 'scrum', 'machine learning', 'ai', 'data science',
            'c++', 'c#', '.net', 'spring', 'hibernate', 'bootstrap', 'jquery', 'php', 'laravel'
        ]
        
        found_skills = []
        text_lower = text_content.lower()
        for skill in skill_keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        
        return {
            "name": potential_name or "Not found",
            "email": emails[0] if emails else "Not found",
            "phone": phones[0] if phones else "Not found",
            "mobile": phones[0] if phones else "Not found",
            "skills": found_skills,
            "languages": languages,
            "raw_text": text_content[:1000] + "..." if len(text_content) > 1000 else text_content,
            "education": education_info,
            "experience": experience_info,
            "parsing_method": "Simple text extraction (fallback)"
        }
        
    except Exception as e:
        print(f"Simple text extraction failed: {e}")
        return None


def extract_education(text_content):
    """Extract education information from resume text"""
    import re
    
    education_info = {
        "college_name": "Not extracted",
        "degree": "Not extracted",
        "details": []
    }
    
    try:
        # Look for education section
        education_patterns = [
            r'EDUCATION\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEXPERIENCE|\nSKILLS|\Z)',
            r'ACADEMIC\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEXPERIENCE|\nSKILLS|\Z)',
            r'QUALIFICATION\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEXPERIENCE|\nSKILLS|\Z)'
        ]
        
        education_text = ""
        for pattern in education_patterns:
            match = re.search(pattern, text_content, re.IGNORECASE | re.DOTALL)
            if match:
                education_text = match.group(1).strip()
                break
        
        if education_text:
            # Extract college/university names
            college_patterns = [
                r'([A-Z][a-zA-Z\s&]+(?:College|University|Institute|School|Academy)[a-zA-Z\s&]*)',
                r'([A-Z][a-zA-Z\s&]+(?:Engineering|Technology|Science|Management)[a-zA-Z\s&]*)'
            ]
            
            colleges = []
            for pattern in college_patterns:
                matches = re.findall(pattern, education_text)
                colleges.extend(matches)
            
            if colleges:
                education_info["college_name"] = colleges[0].strip()
            
            # Extract degrees
            degree_patterns = [
                r'(B\.?Tech|Bachelor|B\.?E\.?|B\.?Sc|M\.?Tech|Master|M\.?E\.?|M\.?Sc|PhD|Doctorate)[a-zA-Z\s]*',
                r'(INFORMATION TECHNOLOGY|Computer Science|Electronics|Mechanical|Civil|Electrical)[a-zA-Z\s]*'
            ]
            
            degrees = []
            for pattern in degree_patterns:
                matches = re.findall(pattern, education_text, re.IGNORECASE)
                degrees.extend(matches)
            
            if degrees:
                education_info["degree"] = degrees[0].strip()
            
            # Extract additional details (CGPA, percentage, etc.)
            lines = education_text.split('\n')
            for line in lines[:10]:  # Check first 10 lines
                line = line.strip()
                if any(keyword in line.lower() for keyword in ['cgpa', 'gpa', 'percentage', '%']):
                    education_info["details"].append(line)
    
    except Exception as e:
        print(f"Error extracting education: {e}")
    
    return education_info


def extract_experience(text_content):
    """Extract experience information from resume text"""
    import re
    
    experience_info = {
        "designation": "Not extracted",
        "company_names": [],
        "total_experience": "0 years",
        "details": []
    }
    
    try:
        # Look for experience section
        experience_patterns = [
            r'EXPERIENCE\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEDUCATION|\nSKILLS|\Z)',
            r'WORK EXPERIENCE\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEDUCATION|\nSKILLS|\Z)',
            r'EMPLOYMENT\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEDUCATION|\nSKILLS|\Z)'
        ]
        
        experience_text = ""
        for pattern in experience_patterns:
            match = re.search(pattern, text_content, re.IGNORECASE | re.DOTALL)
            if match:
                experience_text = match.group(1).strip()
                break
        
        # If no dedicated experience section, look for internships or projects
        if not experience_text:
            internship_patterns = [
                r'INTERNSHIP\s*\n(.*?)(?=\n[A-Z]{2,}|\nPROJECTS|\nEDUCATION|\nSKILLS|\Z)',
                r'PROJECTS\s*\n(.*?)(?=\n[A-Z]{2,}|\nEXPERIENCE|\nEDUCATION|\nSKILLS|\Z)'
            ]
            
            for pattern in internship_patterns:
                match = re.search(pattern, text_content, re.IGNORECASE | re.DOTALL)
                if match:
                    experience_text = match.group(1).strip()
                    experience_info["designation"] = "Fresher/Student Projects"
                    break
        
        if experience_text:
            # Extract company names
            company_patterns = [
                r'([A-Z][a-zA-Z\s&]+(?:Ltd|Limited|Inc|Corp|Company|Technologies|Solutions|Systems)[a-zA-Z\s&]*)',
                r'([A-Z][a-zA-Z\s&]{10,50})'  # General company name pattern
            ]
            
            companies = []
            for pattern in company_patterns:
                matches = re.findall(pattern, experience_text)
                companies.extend([match.strip() for match in matches if len(match.strip()) > 5])
            
            experience_info["company_names"] = list(set(companies))[:3]  # Limit to 3 companies
            
            # Extract years of experience
            year_patterns = [
                r'(\d+)\s*(?:years?|yrs?)',
                r'(\d+)\s*(?:months?|mos?)',
                r'(\d{4})\s*-\s*(\d{4})'  # Year ranges
            ]
            
            for pattern in year_patterns:
                matches = re.findall(pattern, experience_text, re.IGNORECASE)
                if matches:
                    if len(matches[0]) == 2:  # Year range
                        years = int(matches[0][1]) - int(matches[0][0])
                        experience_info["total_experience"] = f"{years} years"
                    else:
                        experience_info["total_experience"] = f"{matches[0]} years"
                    break
        
        # If still no experience found, check if it's a fresh graduate
        if experience_info["designation"] == "Not extracted":
            if any(keyword in text_content.lower() for keyword in ['graduate', 'fresher', 'entry-level', 'seeking', '2024', '2025']):
                experience_info["designation"] = "Fresh Graduate"
                experience_info["total_experience"] = "0 years (Fresher)"
    
    except Exception as e:
        print(f"Error extracting experience: {e}")
    
    return experience_info


def extract_languages(text_content):
    """Extract languages from resume text"""
    import re
    
    languages = []
    
    try:
        # Common languages to look for
        language_keywords = [
            'english', 'hindi', 'tamil', 'telugu', 'kannada', 'malayalam', 'marathi', 'gujarati',
            'bengali', 'punjabi', 'urdu', 'french', 'german', 'spanish', 'japanese', 'chinese',
            'korean', 'arabic', 'russian'
        ]
        
        text_lower = text_content.lower()
        
        # Look for language section first
        language_patterns = [
            r'LANGUAGE\s*\n(.*?)(?=\n[A-Z]{2,}|\nSKILLS|\nEDUCATION|\nEXPERIENCE|\Z)',
            r'LANGUAGES\s*\n(.*?)(?=\n[A-Z]{2,}|\nSKILLS|\nEDUCATION|\nEXPERIENCE|\Z)'
        ]
        
        language_text = ""
        for pattern in language_patterns:
            match = re.search(pattern, text_content, re.IGNORECASE | re.DOTALL)
            if match:
                language_text = match.group(1).strip().lower()
                break
        
        # Search in language section or entire text
        search_text = language_text if language_text else text_lower
        
        for lang in language_keywords:
            if lang in search_text:
                languages.append(lang.capitalize())
        
        # Remove duplicates and limit to reasonable number
        languages = list(set(languages))[:5]
        
        # If no languages found, assume English for Indian resumes
        if not languages and any(keyword in text_lower for keyword in ['india', 'indian', 'tamil nadu', 'mumbai', 'delhi', 'bangalore']):
            languages = ['English']
    
    except Exception as e:
        print(f"Error extracting languages: {e}")
    
    return languages


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    Converts a .docx file to .pdf using the docx2pdf library.
    Returns the path to the new PDF file or an empty string on failure.
    """
    try:
        # Generate output PDF path
        filename = os.path.basename(docx_path)
        pdf_filename = filename.replace('.docx', '.pdf')
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        # Convert using docx2pdf
        convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            return ""
    except Exception as e:
        print(f"Error during conversion: {e}")
        return ""

